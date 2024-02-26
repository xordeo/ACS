import sys
import sqlite3
import datetime as dt

from docx import Document
from random import choice
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QLabel, QLineEdit, QPushButton,
                             QTableWidget, QTableWidgetItem, QComboBox, QAbstractItemView, QHeaderView, QFileDialog)
from PyQt5.QtGui import QIcon, QFont
from PyQt5.QtCore import Qt


class LoginWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        self.setGeometry(700, 350, 400, 400)
        self.setFixedSize(400, 400)
        self.setWindowIcon(QIcon("icons/login_icon.png"))
        self.setWindowTitle("Вход в учётную запись СКУД")

        self.loginLabel = QLabel("Введите логин:", self)
        self.loginLabel.setFont(QFont("Calibri", 15))
        self.loginLabel.move(110, 100)
        self.loginLabel.resize(200, 25)

        self.loginEdit = QLineEdit("", self)
        self.loginEdit.move(110, 130)
        self.loginEdit.resize(175, 25)

        self.passwordLabel = QLabel("Введите пароль:", self)
        self.passwordLabel.setFont(QFont("Calibri", 15))
        self.passwordLabel.move(110, 180)
        self.passwordLabel.resize(200, 25)

        self.passwordEdit = QLineEdit("", self)
        self.passwordEdit.move(110, 210)
        self.passwordEdit.resize(175, 25)

        self.messageLabel = QLabel("", self)
        self.messageLabel.setFont(QFont("Calibri", 10))
        self.messageLabel.setStyleSheet("QLabel { color : red; }")
        self.messageLabel.move(65, 250)
        self.messageLabel.resize(300, 30)

        self.loginButton = QPushButton("Войти", self)
        self.loginButton.setFont(QFont("Calibri", 12))
        self.loginButton.move(145, 290)
        self.loginButton.resize(100, 30)
        self.loginButton.clicked.connect(self.login_in)

    def login_in(self):
        self.connection = sqlite3.connect("ACS.db")
        self.cursor = self.connection.cursor()

        # поиск в базе данных введённой пары логин-пароль
        self.user_account = self.cursor.execute(f"""SELECT * FROM maintable 
        WHERE employee_login = '{self.loginEdit.text()}' 
        AND employee_password = '{self.passwordEdit.text()}'""").fetchone()

        self.connection.close()

        # если учётная запись существует, в зависимости от должности открывается конкретная форма
        if self.user_account:
            self.hide()
            if self.user_account[3] == "Системный администратор":
                self.user_form = SystemAdmin(self, self.user_account)
            elif self.user_account[3] == "Администрация":
                self.user_form = Administration(self, self.user_account)
            else:
                self.user_form = Employee(self, self.user_account)
            self.user_form.show()
        else:
            self.messageLabel.setText("Ошибка входа! Пара логин-пароль не найдена. ❌")


class SystemAdmin(QWidget):
    def __init__(self, *args):
        super().__init__()
        self.initUI(args)

    def initUI(self, args):
        self.setGeometry(300, 200, 980, 500)
        self.setFixedSize(980, 500)
        self.setWindowIcon(QIcon("icons/form_icon.png"))
        self.setWindowTitle("СКУД")

        self.connection = sqlite3.connect("ACS.db")
        self.cursor = self.connection.cursor()

        self.user_data = args[-1]
        self.ID = self.user_data[0]
        self.arrival_time = self.user_data[5][:6]
        self.care_time = self.user_data[6][:6]
        self.time_now = dt.datetime.now().strftime("%d %b")

        self.infoLabel = QLabel(f"👤 ФИО: {self.user_data[4]} | Должность: {self.user_data[3]}", self)
        self.infoLabel.setFont(QFont("Calibri", 12))
        self.infoLabel.move(20, 15)
        self.infoLabel.resize(700, 25)

        self.tableWidget = QTableWidget(self)
        self.tableWidget.move(0, 50)
        self.tableWidget.resize(750, 400)

        # отключение изменения таблицы напрямую
        self.tableWidget.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.tableWidget.horizontalHeader().setSectionResizeMode(QHeaderView.Fixed)

        self.update_table()

        self.registerButton = QPushButton("Регистрация сотрудника", self)
        self.registerButton.setFont(QFont("Calibri", 11))
        self.registerButton.move(780, 300)
        self.registerButton.resize(170, 35)
        self.registerButton.clicked.connect(self.register_user_account)

        self.changeButton = QPushButton("Изменить данные", self)
        self.changeButton.setFont(QFont("Calibri", 12))
        self.changeButton.move(780, 350)
        self.changeButton.resize(170, 35)
        self.changeButton.clicked.connect(self.change_table)

        self.deleteButton = QPushButton("Удалить сотрудника", self)
        self.deleteButton.setFont(QFont("Calibri", 12))
        self.deleteButton.move(780, 400)
        self.deleteButton.resize(170, 35)
        self.deleteButton.clicked.connect(self.delete_user_account)

    def update_table(self):
        self.tableWidget.clear()
        table = self.cursor.execute(f"""SELECT * FROM maintable""").fetchall()
        self.tableWidget.setColumnCount(7)
        self.tableWidget.setHorizontalHeaderLabels(["ID", "Логин", "Пароль", "Должность", "ФИО",
                                                    "Время прихода", "Время ухода"])
        self.tableWidget.setRowCount(0)
        for i, row in enumerate(table):
            self.tableWidget.setRowCount(
                self.tableWidget.rowCount() + 1)
            for j, elem in enumerate(row):
                self.tableWidget.setItem(
                    i, j, QTableWidgetItem(str(elem)))
        self.tableWidget.resizeColumnsToContents()

    def register_user_account(self):
        self.register_form = RegisterWindow(self)
        self.register_form.show()

    def change_table(self):
        self.change_form = ChangeWindow(self)
        self.change_form.show()

    def delete_user_account(self):
        self.delete_form = DeleteWindow(self)
        self.delete_form.show()


def generate_password():
    password = ""
    for i in range(12):
        password += choice("&$#abcdefghijklnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ1234567890")
    return password


class RegisterWindow(QWidget):
    def __init__(self, main_window):
        super().__init__()
        self.main_window = main_window
        self.initUI()

    def initUI(self):
        self.setGeometry(700, 300, 470, 320)
        self.setFixedSize(470, 320)
        self.setWindowIcon(QIcon("icons/register_icon.png"))
        self.setWindowTitle("Регистрация сотрудника в СКУД")

        self.connection = sqlite3.connect("ACS.db")
        self.cursor = self.connection.cursor()

        self.loginLabel = QLabel("Логин:", self)
        self.loginLabel.setFont(QFont("Calibri", 12))
        self.loginLabel.move(40, 20)
        self.loginLabel.resize(200, 25)

        self.loginEdit = QLineEdit(self)
        self.loginEdit.move(40, 50)
        self.loginEdit.resize(200, 25)

        self.passwordLabel = QLabel("Пароль:", self)
        self.passwordLabel.setFont(QFont("Calibri", 12))
        self.passwordLabel.move(40, 80)
        self.passwordLabel.resize(200, 25)

        self.passwordEdit = QLineEdit(self)
        self.passwordEdit.move(40, 110)
        self.passwordEdit.resize(200, 25)

        self.passwordButton = QPushButton("Сгенерировать пароль", self)
        self.passwordButton.setFont(QFont("Calibri", 12))
        self.passwordButton.move(260, 109)
        self.passwordButton.resize(180, 27)
        self.passwordButton.clicked.connect(self.set_password)

        self.job_titleLabel = QLabel("Должность:", self)
        self.job_titleLabel.setFont(QFont("Calibri", 12))
        self.job_titleLabel.move(40, 140)
        self.job_titleLabel.resize(200, 25)

        self.job_titleEdit = QLineEdit(self)
        self.job_titleEdit.move(40, 170)
        self.job_titleEdit.resize(200, 25)

        self.nameLabel = QLabel("ФИО:", self)
        self.nameLabel.setFont(QFont("Calibri", 12))
        self.nameLabel.move(40, 210)
        self.nameLabel.resize(200, 25)

        self.nameEdit = QLineEdit(self)
        self.nameEdit.move(40, 240)
        self.nameEdit.resize(200, 25)

        self.messageLabel = QLabel("", self)
        self.messageLabel.setFont(QFont("Calibri", 10))
        self.messageLabel.move(260, 209)
        self.messageLabel.resize(300, 30)

        self.registerButton = QPushButton("Зарегистрировать", self)
        self.registerButton.setFont(QFont("Calibri", 12))
        self.registerButton.move(260, 239)
        self.registerButton.resize(180, 27)
        self.registerButton.clicked.connect(self.register_user_account)

    def set_password(self):
        self.passwordEdit.setText(generate_password())

    def register_user_account(self):
        if (self.loginEdit.text() != "" and self.passwordEdit.text() != ""
                and self.job_titleEdit.text() != "" and self.nameEdit.text()):
            if (self.loginEdit.text() in
                    [login[0] for login in self.cursor.execute("""SELECT employee_login FROM maintable""").fetchall()]):
                self.messageLabel.setStyleSheet("QLabel { color : red; }")
                self.messageLabel.setText("Такой логин уже существует!")
            else:
                time_now = dt.datetime.now().strftime("%d %b %H:%M")
                self.cursor.execute(f"""INSERT INTO 
                maintable(employee_login, employee_password, job_title, name, arrival_time, care_time) 
                VALUES('{self.loginEdit.text()}', '{self.passwordEdit.text()}', 
                '{self.job_titleEdit.text()}', '{self.nameEdit.text()}', '{time_now}', '{time_now}')""")
                self.connection.commit()
                self.main_window.update_table()
                self.messageLabel.setStyleSheet("QLabel { color : green; }")
                self.messageLabel.setText("Регистрация прошла успешно!")
        else:
            self.messageLabel.setStyleSheet("QLabel { color : red; }")
            self.messageLabel.setText("Указаны не все данные!")


class ChangeWindow(QWidget):
    def __init__(self, main_window):
        super().__init__()
        self.main_window = main_window
        self.initUI()

    def initUI(self):
        self.setGeometry(700, 300, 450, 500)
        self.setFixedSize(450, 500)
        self.setWindowIcon(QIcon("icons/change_icon.png"))
        self.setWindowTitle("Изменить данные СКУД")

        self.connection = sqlite3.connect("ACS.db")
        self.cursor = self.connection.cursor()

        self.IDLabel = QLabel("Выберите ID сотрудника:", self)
        self.IDLabel.setFont(QFont("Calibri", 12))
        self.IDLabel.move(40, 20)
        self.IDLabel.resize(200, 25)

        self.IDComboBox = QComboBox(self)
        self.IDComboBox.move(40, 50)
        self.IDComboBox.resize(200, 25)
        self.IDComboBox.addItems(
            [str(elem[0]) for elem in self.cursor.execute("""SELECT id FROM maintable""").fetchall()])
        self.IDComboBox.currentTextChanged.connect(self.check_user_data)

        self.loginLabel = QLabel("Логин:", self)
        self.loginLabel.setFont(QFont("Calibri", 12))
        self.loginLabel.move(40, 80)
        self.loginLabel.resize(200, 25)

        self.loginEdit = QLineEdit(self)
        self.loginEdit.move(40, 110)
        self.loginEdit.resize(200, 25)

        self.passwordLabel = QLabel("Пароль:", self)
        self.passwordLabel.setFont(QFont("Calibri", 12))
        self.passwordLabel.move(40, 140)
        self.passwordLabel.resize(200, 25)

        self.passwordEdit = QLineEdit(self)
        self.passwordEdit.move(40, 170)
        self.passwordEdit.resize(200, 25)

        self.passwordButton = QPushButton("Сгенерировать пароль", self)
        self.passwordButton.setFont(QFont("Calibri", 12))
        self.passwordButton.move(260, 169)
        self.passwordButton.resize(180, 27)
        self.passwordButton.clicked.connect(self.set_password)

        self.job_titleLabel = QLabel("Должность:", self)
        self.job_titleLabel.setFont(QFont("Calibri", 12))
        self.job_titleLabel.move(40, 200)
        self.job_titleLabel.resize(200, 25)

        self.job_titleEdit = QLineEdit(self)
        self.job_titleEdit.move(40, 230)
        self.job_titleEdit.resize(200, 25)

        self.nameLabel = QLabel("ФИО:", self)
        self.nameLabel.setFont(QFont("Calibri", 12))
        self.nameLabel.move(40, 260)
        self.nameLabel.resize(200, 25)

        self.nameEdit = QLineEdit(self)
        self.nameEdit.move(40, 290)
        self.nameEdit.resize(200, 25)

        self.arrival_timeLabel = QLabel("Время прихода:", self)
        self.arrival_timeLabel.setFont(QFont("Calibri", 12))
        self.arrival_timeLabel.move(40, 320)
        self.arrival_timeLabel.resize(200, 25)

        self.arrival_timeEdit = QLineEdit(self)
        self.arrival_timeEdit.move(40, 350)
        self.arrival_timeEdit.resize(200, 25)

        self.care_timeLabel = QLabel("Время ухода:", self)
        self.care_timeLabel.setFont(QFont("Calibri", 12))
        self.care_timeLabel.move(40, 380)
        self.care_timeLabel.resize(200, 25)

        self.care_timeEdit = QLineEdit(self)
        self.care_timeEdit.move(40, 410)
        self.care_timeEdit.resize(200, 25)

        self.saveButton = QPushButton("Сохранить изменения", self)
        self.saveButton.setFont(QFont("Calibri", 12))
        self.saveButton.move(260, 289)
        self.saveButton.resize(180, 27)
        self.saveButton.clicked.connect(self.save_changes)

        self.messageLabel = QLabel("", self)
        self.messageLabel.setFont(QFont("Calibri", 10))
        self.messageLabel.resize(200, 25)

        self.check_user_data()

    def set_password(self):
        self.passwordEdit.setText(generate_password())

    def save_changes(self):
        ID = self.IDComboBox.currentText()
        logins = [elem[0] for elem in self.cursor.execute(
            f"""SELECT employee_login FROM maintable WHERE id != {ID}""").fetchall()]
        if self.loginEdit.text() not in logins:
            self.connection.execute(
                f"""UPDATE maintable SET employee_login = '{self.loginEdit.text()}' WHERE id = {ID}""")
            self.connection.execute(
                f"""UPDATE maintable SET employee_password = '{self.passwordEdit.text()}' WHERE id = {ID}""")
            self.connection.execute(
                f"""UPDATE maintable SET job_title = '{self.job_titleEdit.text()}' WHERE id = {ID}""")
            self.connection.execute(
                f"""UPDATE maintable SET name = '{self.nameEdit.text()}' WHERE id = {ID}""")
            self.connection.execute(
                f"""UPDATE maintable SET arrival_time = '{self.arrival_timeEdit.text()}' WHERE id = {ID}""")
            self.connection.execute(
                f"""UPDATE maintable SET care_time = '{self.care_timeEdit.text()}' WHERE id = {ID}""")
            self.connection.commit()
            self.messageLabel.setStyleSheet("QLabel { color : green; }")
            self.messageLabel.move(290, 329)
            self.messageLabel.setText("Успешно сохранено!")
            self.main_window.update_table()
        else:
            self.messageLabel.setStyleSheet("QLabel { color : red; }")
            self.messageLabel.move(270, 329)
            self.messageLabel.setText("Такой логин уже существует!")

    def check_user_data(self):
        user_data = self.connection.execute(
            f"""SELECT * FROM maintable WHERE id = {self.IDComboBox.currentText()}""").fetchone()
        self.loginEdit.setText(user_data[1])
        self.passwordEdit.setText(user_data[2])
        self.job_titleEdit.setText(user_data[3])
        self.nameEdit.setText(user_data[4])
        self.arrival_timeEdit.setText(user_data[5])
        self.care_timeEdit.setText(user_data[6])


class DeleteWindow(QWidget):
    def __init__(self, main_window):
        super().__init__()
        self.main_window = main_window
        self.initUI()

    def initUI(self):
        self.setGeometry(700, 300, 450, 200)
        self.setFixedSize(450, 200)
        self.setWindowIcon(QIcon("icons/delete_icon.png"))
        self.setWindowTitle("Удаление сотрудника из СКУД")

        self.IDLabel = QLabel("Выберите ID сотрудника:", self)
        self.IDLabel.setFont(QFont("Calibri", 12))
        self.IDLabel.move(40, 20)
        self.IDLabel.resize(200, 25)

        self.connection = sqlite3.connect("ACS.db")
        self.cursor = self.connection.cursor()

        self.IDComboBox = QComboBox(self)
        self.IDComboBox.move(40, 50)
        self.IDComboBox.resize(200, 25)
        self.IDComboBox.addItems(
            [str(elem[0]) for elem in self.cursor.execute("""SELECT id FROM maintable""").fetchall()])
        self.IDComboBox.currentTextChanged.connect(self.check_user_data)

        self.job_titleLabel = QLabel(self)
        self.job_titleLabel.setFont(QFont("Calibri", 12))
        self.job_titleLabel.move(40, 80)
        self.job_titleLabel.resize(600, 30)

        self.nameLabel = QLabel(self)
        self.nameLabel.setFont(QFont("Calibri", 12))
        self.nameLabel.move(40, 110)
        self.nameLabel.resize(600, 30)

        self.deleteButton = QPushButton("Удалить этого сотрудника", self)
        self.deleteButton.setFont(QFont("Calibri", 12))
        self.deleteButton.move(40, 150)
        self.deleteButton.resize(200, 30)
        self.deleteButton.clicked.connect(self.delete_user_data)

        self.check_user_data()

    def check_user_data(self):
        data = self.connection.execute(
            f"""SELECT job_title, name FROM maintable WHERE id = {self.IDComboBox.currentText()}""").fetchall()[0]
        self.job_titleLabel.setText(f"Должность: {data[0]}")
        self.nameLabel.setText(f"ФИО: {data[1]}")

    def delete_user_data(self):
        self.cursor.execute(f"""DELETE from maintable WHERE id = {self.IDComboBox.currentText()}""")
        self.connection.commit()
        self.main_window.update_table()
        self.hide()


class Administration(QWidget):
    def __init__(self, *args):
        super().__init__()
        self.initUI(args)

    def initUI(self, args):
        self.setGeometry(300, 200, 900, 500)
        self.setFixedSize(900, 500)
        self.setWindowIcon(QIcon("icons/form_icon.png"))
        self.setWindowTitle("СКУД")

        self.connection = sqlite3.connect("ACS.db")
        self.cursor = self.connection.cursor()

        self.user_data = args[-1]
        self.ID = self.user_data[0]
        self.arrival_time = self.user_data[5][:6]
        self.care_time = self.user_data[6][:6]
        self.time_now = dt.datetime.now().strftime("%d %b")

        self.infoLabel = QLabel(f"👤 ФИО: {self.user_data[4]} | Должность: {self.user_data[3]}", self)
        self.infoLabel.setFont(QFont("Calibri", 12))
        self.infoLabel.move(20, 15)
        self.infoLabel.resize(700, 25)

        self.tableWidget = QTableWidget(self)
        self.tableWidget.move(0, 50)
        self.tableWidget.resize(650, 400)

        # отключение изменения таблицы напрямую
        self.tableWidget.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.tableWidget.horizontalHeader().setSectionResizeMode(QHeaderView.Fixed)

        self.filterLabel = QLabel("Фильтрация таблицы:", self)
        self.filterLabel.setFont(QFont("Calibri", 12))
        self.filterLabel.move(670, 15)
        self.filterLabel.resize(700, 25)

        self.filterComboBox = QComboBox(self)
        self.filterComboBox.setFont(QFont("Calibri", 10))
        self.filterComboBox.move(670, 50)
        self.filterComboBox.resize(200, 30)
        self.filterComboBox.addItems(["ID по возрастанию", "ID по убыванию",
                                      "Должность А-Я", "Должность Я-А",
                                      "ФИО А-Я", "ФИО Я-А",
                                      "Время прихода по возрастанию", "Время прихода по убыванию",
                                      "Время ухода по возрастанию", "Время ухода по убыванию"])
        self.filterComboBox.currentTextChanged.connect(self.filter_table)

        self.reportButton = QPushButton("Сохранить отчёт", self)
        self.reportButton.setFont(QFont("Calibri", 12))
        self.reportButton.move(680, 400)
        self.reportButton.resize(170, 35)
        self.reportButton.clicked.connect(self.create_report)

        self.update_table()

    def update_table(self):
        self.tableWidget.clear()
        table = self.cursor.execute(
            f"""SELECT id, job_title, name, arrival_time, care_time FROM maintable""").fetchall()
        self.tableWidget.setColumnCount(5)
        self.tableWidget.setHorizontalHeaderLabels(["ID", "Должность", "ФИО",
                                                    "Время прихода", "Время ухода"])
        self.tableWidget.setRowCount(0)
        for i, row in enumerate(table):
            self.tableWidget.setRowCount(
                self.tableWidget.rowCount() + 1)
            for j, elem in enumerate(row):
                self.tableWidget.setItem(
                    i, j, QTableWidgetItem(str(elem)))
        self.tableWidget.resizeColumnsToContents()

    def filter_table(self):
        filter = self.filterComboBox.currentText()
        if filter == "ID по возрастанию":
            self.tableWidget.sortItems(0, Qt.AscendingOrder)
        if filter == "ID по убыванию":
            self.tableWidget.sortItems(0, Qt.DescendingOrder)
        if filter == "Должность А-Я":
            self.tableWidget.sortItems(1, Qt.AscendingOrder)
        if filter == "Должность Я-А":
            self.tableWidget.sortItems(1, Qt.DescendingOrder)
        if filter == "ФИО А-Я":
            self.tableWidget.sortItems(2, Qt.AscendingOrder)
        if filter == "ФИО Я-А":
            self.tableWidget.sortItems(2, Qt.DescendingOrder)
        if filter == "Время прихода по возрастанию":
            self.tableWidget.sortItems(3, Qt.AscendingOrder)
        if filter == "Время прихода по убыванию":
            self.tableWidget.sortItems(3, Qt.DescendingOrder)
        if filter == "Время ухода по возрастанию":
            self.tableWidget.sortItems(4, Qt.AscendingOrder)
        if filter == "Время ухода по убыванию":
            self.tableWidget.sortItems(4, Qt.DescendingOrder)

    def create_report(self):
        directory = list(QFileDialog.getSaveFileName(
            self, 'Выбрать путь сохранения отчёта', 'report.docx', 'Документ (*.docx)'))[0]
        if directory:
            today_date = dt.datetime.now().strftime("%d.%m.%Y")
            doc = Document()
            doc.add_heading(f"Отчёт о посещаемости сотрудников на {today_date}", 0)
            table = doc.add_table(rows=self.tableWidget.rowCount() + 1, cols=self.tableWidget.columnCount())
            table.style = "Table Grid"
            sql_table = self.cursor.execute(
                f"""SELECT id, job_title, name, arrival_time, care_time FROM maintable""").fetchall()
            for text in enumerate(["ID", "Должность", "ФИО", "Время прихода", "Время ухода"]):
                table.cell(0, text[0]).text = text[1]

            for i, row in enumerate(sql_table):
                for j, column in enumerate(row):
                    table.cell(i + 1, j).text = str(column)
            doc.save(directory)


class Employee(QWidget):
    def __init__(self, *args):
        super().__init__()
        self.initUI(args)

    def initUI(self, args):
        self.setGeometry(300, 300, 700, 230)
        self.setFixedSize(700, 230)
        self.setWindowIcon(QIcon("icons/form_icon.png"))
        self.setWindowTitle("СКУД")

        self.connection = sqlite3.connect("ACS.db")
        self.cursor = self.connection.cursor()

        self.user_data = args[-1]
        self.ID = self.user_data[0]
        self.arrival_time = self.user_data[5][:6]
        self.care_time = self.user_data[6][:6]
        self.time_now = dt.datetime.now().strftime("%d %b")

        self.infoLabel = QLabel(f"👤 ФИО: {self.user_data[4]} | Должность: {self.user_data[3]}", self)
        self.infoLabel.setFont(QFont("Calibri", 12))
        self.infoLabel.move(40, 20)
        self.infoLabel.resize(700, 25)

        self.arrival_timeLabel = QLabel("Отметка о приходе:", self)
        self.arrival_timeLabel.setFont(QFont("Calibri", 12))
        self.arrival_timeLabel.move(40, 50)
        self.arrival_timeLabel.resize(200, 25)

        self.arrival_timeEdit = QLineEdit(self)
        self.arrival_timeEdit.move(40, 80)
        self.arrival_timeEdit.resize(200, 25)
        self.arrival_timeEdit.setDisabled(True)

        self.check_inButton1 = QPushButton("Отметиться", self)
        self.check_inButton1.setFont(QFont("Calibri", 12))
        self.check_inButton1.move(85, 130)
        self.check_inButton1.resize(100, 50)
        self.check_inButton1.clicked.connect(self.check_in1)

        self.care_timeLabel = QLabel("Отметка об уходе:", self)
        self.care_timeLabel.setFont(QFont("Calibri", 12))
        self.care_timeLabel.move(340, 50)
        self.care_timeLabel.resize(200, 25)

        self.care_timeEdit = QLineEdit(self)
        self.care_timeEdit.move(340, 80)
        self.care_timeEdit.resize(200, 25)
        self.care_timeEdit.setDisabled(True)

        self.check_inButton2 = QPushButton("Отметиться", self)
        self.check_inButton2.setFont(QFont("Calibri", 12))
        self.check_inButton2.move(385, 130)
        self.check_inButton2.resize(100, 50)
        self.check_inButton2.clicked.connect(self.check_in2)

        self.login_update()

    def check_in1(self):
        self.check_inButton1.setDisabled(True)
        self.check_inButton2.setDisabled(False)
        check_in = dt.datetime.now().strftime("%d %b %H:%M")
        self.arrival_timeEdit.setText(check_in)
        self.cursor.execute(f"""UPDATE maintable SET arrival_time = '{check_in}' WHERE id = '{self.ID}'""")
        self.connection.commit()

    def check_in2(self):
        self.check_inButton2.setDisabled(True)
        check_in = dt.datetime.now().strftime("%d %b %H:%M")
        self.care_timeEdit.setText(check_in)
        self.cursor.execute(f"""UPDATE maintable SET care_time = '{check_in}' WHERE id = '{self.ID}'""")
        self.connection.commit()

    def login_update(self):
        if self.arrival_time == self.time_now:
            self.arrival_timeEdit.setText(
                self.cursor.execute(f"""SELECT arrival_time from maintable WHERE id = '{self.ID}'""").fetchone()[0])
            self.check_inButton1.setDisabled(True)
        if self.care_time == self.time_now:
            self.care_timeEdit.setText(
                self.cursor.execute(f"""SELECT care_time from maintable WHERE id = '{self.ID}'""").fetchone()[0])
            self.check_inButton2.setDisabled(True)
        if self.care_timeEdit.text() == "":
            self.check_inButton2.setDisabled(True if self.check_inButton1.isEnabled() else False)


if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = LoginWindow()
    window.show()
    sys.exit(app.exec())
